#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文本提取模块

支持从 PDF、Word、Excel、TXT、CSV、Markdown、WPS 格式文件中提取文本内容。
"""

import os
import csv
from typing import Dict, Any, Optional


def extract_text_from_file(file_path: str, enable_ocr: bool = False) -> Dict[str, Any]:
    """
    从文件中提取文本内容

    Args:
        file_path: 文件路径
        enable_ocr: 是否启用 OCR 支持（用于扫描件 PDF）

    Returns:
        包含提取结果的字典，格式为：
        {
            "success": bool,
            "data": {"text": str},
            "error": str
        }
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return {
                "success": False,
                "data": {},
                "error": f"文件不存在: {file_path}"
            }

        # 检查文件权限
        if not os.access(file_path, os.R_OK):
            return {
                "success": False,
                "data": {},
                "error": f"权限不足，无法读取文件: {file_path}"
            }

        # 智能格式检测
        file_type = _detect_file_type(file_path)
        
        # 根据文件类型调用不同的提取函数
        if file_type == "txt":
            text = _extract_from_txt(file_path)
        elif file_type == "pdf":
            if enable_ocr:
                text = _extract_from_pdf_with_ocr(file_path)
            else:
                text = _extract_from_pdf(file_path)
        elif file_type == "docx":
            text = _extract_from_docx(file_path)
        elif file_type in ["xlsx", "xls"]:
            text = _extract_from_excel(file_path)
        elif file_type == "csv":
            text = _extract_from_csv(file_path)
        elif file_type == "md":
            text = _extract_from_markdown(file_path)
        elif file_type == "wps":
            text = _extract_from_wps(file_path)
        elif file_type == "et":
            text = _extract_from_wps_et(file_path)
        else:
            return {
                "success": False,
                "data": {},
                "error": f"不支持的文件格式: {os.path.splitext(file_path)[1].lower()}"
            }

        return {
            "success": True,
            "data": {"text": text},
            "error": ""
        }

    except Exception as e:
        return {
            "success": False,
            "data": {},
            "error": f"提取文本失败: {str(e)}"
        }


def _detect_file_type(file_path: str) -> str:
    """
    智能检测文件类型

    Args:
        file_path: 文件路径

    Returns:
        文件类型字符串
    """
    # 首先根据文件扩展名判断
    file_ext = os.path.splitext(file_path)[1].lower()
    
    ext_map = {
        ".txt": "txt",
        ".pdf": "pdf",
        ".docx": "docx",
        ".xlsx": "xlsx",
        ".xls": "xls",
        ".csv": "csv",
        ".md": "md",
        ".markdown": "md",
        ".wps": "wps",
        ".et": "et"
    }
    
    if file_ext in ext_map:
        return ext_map[file_ext]
    
    # 尝试根据文件内容判断（简单实现）
    try:
        with open(file_path, 'rb') as f:
            header = f.read(100)
        
        # PDF 文件头
        if header.startswith(b'%PDF'):
            return "pdf"
        # ZIP 格式文件（docx, xlsx, wps, et）
        elif header.startswith(b'PK\x03\x04'):
            return "docx"  # 默认返回 docx，实际使用时会根据扩展名更准确判断
        # CSV 文件（尝试读取）
        elif b',' in header or b';' in header:
            return "csv"
    except:
        pass
    
    return "unsupported"  # 默认返回 unsupported


def _extract_from_txt(file_path: str) -> str:
    """
    从 TXT 文件中提取文本

    Args:
        file_path: TXT 文件路径

    Returns:
        提取的文本内容
    """
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def _extract_from_pdf(file_path: str) -> str:
    """
    从 PDF 文件中提取文本

    Args:
        file_path: PDF 文件路径

    Returns:
        提取的文本内容
    """
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text()
        return text
    except ImportError:
        raise Exception("PyPDF2 库未安装，请运行 'pip install PyPDF2'")


def _extract_from_pdf_with_ocr(file_path: str) -> str:
    """
    使用 OCR 从 PDF 扫描件中提取文本

    Args:
        file_path: PDF 文件路径

    Returns:
        提取的文本内容
    """
    try:
        import pytesseract
        from PIL import Image
        import pdf2image
        
        # 将 PDF 转换为图像
        images = pdf2image.convert_from_path(file_path)
        text = ""
        
        # 对每个图像进行 OCR
        for image in images:
            text += pytesseract.image_to_string(image, lang='chi_sim+eng') + '\n'
        
        return text
    except ImportError:
        raise Exception("OCR 依赖库未安装，请运行 'pip install pytesseract pdf2image Pillow'")
    except Exception as e:
        # OCR 失败时回退到普通 PDF 提取
        return _extract_from_pdf(file_path)


def _extract_from_docx(file_path: str) -> str:
    """
    从 Word 文件中提取文本

    Args:
        file_path: Word 文件路径

    Returns:
        提取的文本内容
    """
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        # 提取段落文本
        for para in doc.paragraphs:
            text += para.text + '\n'
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join([cell.text for cell in row.cells])
                text += row_text + '\n'
        return text
    except ImportError:
        raise Exception("python-docx 库未安装，请运行 'pip install python-docx'")


def _extract_from_excel(file_path: str) -> str:
    """
    从 Excel 文件中提取文本

    Args:
        file_path: Excel 文件路径

    Returns:
        提取的文本内容
    """
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path)
        text = ""
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"工作表: {sheet_name}\n"
            # 处理合并单元格
            merged_cells = sheet.merged_cells.ranges
            merged_dict = {}
            for merged in merged_cells:
                top_left = merged.min_row, merged.min_col
                value = sheet.cell(row=top_left[0], column=top_left[1]).value
                for row in range(merged.min_row, merged.max_row + 1):
                    for col in range(merged.min_col, merged.max_col + 1):
                        merged_dict[(row, col)] = value
            
            # 提取单元格内容
            for row in range(1, sheet.max_row + 1):
                row_cells = []
                for col in range(1, sheet.max_column + 1):
                    if (row, col) in merged_dict:
                        row_cells.append(str(merged_dict[(row, col)]) if merged_dict[(row, col)] is not None else '')
                    else:
                        cell_value = sheet.cell(row=row, column=col).value
                        row_cells.append(str(cell_value) if cell_value is not None else '')
                row_text = '\t'.join(row_cells)
                text += row_text + '\n'
            text += '\n'
        return text
    except ImportError:
        raise Exception("openpyxl 库未安装，请运行 'pip install openpyxl'")


def _extract_from_csv(file_path: str) -> str:
    """
    从 CSV 文件中提取文本

    Args:
        file_path: CSV 文件路径

    Returns:
        提取的文本内容
    """
    text = ""
    try:
        # 尝试使用 utf-8 编码
        with open(file_path, 'r', encoding='utf-8', newline='') as f:
            reader = csv.reader(f)
            for row in reader:
                row_text = '\t'.join(row)
                text += row_text + '\n'
    except UnicodeDecodeError:
        # 尝试使用 gbk 编码
        with open(file_path, 'r', encoding='gbk', newline='', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                row_text = '\t'.join(row)
                text += row_text + '\n'
    return text


def _extract_from_markdown(file_path: str) -> str:
    """
    从 Markdown 文件中提取文本

    Args:
        file_path: Markdown 文件路径

    Returns:
        提取的文本内容
    """
    try:
        import markdown
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            md_content = f.read()
        
        # 将 Markdown 转换为 HTML
        html = markdown.markdown(md_content)
        
        # 使用 BeautifulSoup 提取纯文本
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n')
        
        return text
    except ImportError:
        # 如果没有安装依赖，直接读取文件内容
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


def _extract_from_wps(file_path: str) -> str:
    """
    从 WPS 文件中提取文本

    Args:
        file_path: WPS 文件路径

    Returns:
        提取的文本内容
    """
    try:
        # WPS 文件本质是 ZIP 格式，尝试使用 python-docx 读取
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + '\n'
        return text
    except Exception:
        # 如果失败，尝试直接读取文件内容
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


def _extract_from_wps_et(file_path: str) -> str:
    """
    从 WPS 表格文件中提取文本

    Args:
        file_path: WPS 表格文件路径

    Returns:
        提取的文本内容
    """
    try:
        # WPS 表格文件本质是 ZIP 格式，尝试使用 openpyxl 读取
        from openpyxl import load_workbook
        wb = load_workbook(file_path)
        text = ""
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"工作表: {sheet_name}\n"
            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                text += row_text + '\n'
            text += '\n'
        return text
    except Exception:
        # 如果失败，尝试直接读取文件内容
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
